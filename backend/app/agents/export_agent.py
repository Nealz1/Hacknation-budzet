from sqlalchemy.orm import Session
from sqlalchemy import func
from typing import List, Dict, Optional
from datetime import datetime, date
from io import BytesIO
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from ..models import BudgetEntry, Department, GlobalLimit, PriorityLevel

class ExportAgent:
    
    def __init__(self, db: Session):
        self.db = db
    
    def export_budget_to_excel(self, year: int = 2025, 
                               department_code: Optional[str] = None) -> BytesIO:
        query = self.db.query(BudgetEntry)
        
        if department_code:
            query = query.join(Department).filter(Department.code == department_code)
        
        entries = query.all()
        
        data = []
        for entry in entries:
            data.append({
                'ID': entry.id,
                'Departament': entry.department.code if entry.department else 'N/A',
                'Paragraf': entry.paragraf,
                'Nazwa zadania': entry.nazwa_zadania or entry.opis_projektu or '',
                'Kwota 2025': entry.kwota_2025 or 0,
                'Kwota 2026': entry.kwota_2026 or 0,
                'Kwota 2027': entry.kwota_2027 or 0,
                'Kwota 2028': entry.kwota_2028 or 0,
                'Kwota 2029': entry.kwota_2029 or 0,
                'Suma': sum([
                    entry.kwota_2025 or 0,
                    entry.kwota_2026 or 0,
                    entry.kwota_2027 or 0,
                    entry.kwota_2028 or 0,
                    entry.kwota_2029 or 0
                ]),
                'Priorytet': entry.priority if entry.priority else 'medium',
                'Obligatoryjne': 'TAK' if entry.is_obligatory else 'NIE',
                'Status': entry.status if entry.status else 'draft',
                'Źródło fin.': entry.zrodlo_finansowania or '',
                'BZ': entry.beneficjent_zadaniowy or '',
                'Uzasadnienie': (entry.szczegolowe_uzasadnienie or '')[:200],
                'Uwagi': entry.uwagi or '',
                'Zwalidowane': 'TAK' if entry.compliance_validated else 'NIE'
            })
        
        df = pd.DataFrame(data)
        
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Pozycje budżetowe', index=False)
            
            summary_data = {
                'Metryka': [
                    'Liczba pozycji',
                    f'Suma {year}',
                    'Pozycje obligatoryjne',
                    'Suma obligatoryjna',
                    'Limit globalny',
                    'Różnica'
                ],
                'Wartość': [
                    len(entries),
                    df['Kwota 2025'].sum() if 'Kwota 2025' in df.columns else 0,
                    len([e for e in entries if e.is_obligatory]),
                    sum(e.kwota_2025 or 0 for e in entries if e.is_obligatory),
                    self._get_global_limit(year),
                    df['Kwota 2025'].sum() - self._get_global_limit(year) if 'Kwota 2025' in df.columns else 0
                ]
            }
            pd.DataFrame(summary_data).to_excel(writer, sheet_name='Podsumowanie', index=False)
            
            if not department_code:
                dept_data = []
                for dept in self.db.query(Department).all():
                    dept_entries = [e for e in entries if e.department_id == dept.id]
                    dept_total = sum(e.kwota_2025 or 0 for e in dept_entries)
                    dept_data.append({
                        'Departament': dept.code,
                        'Nazwa': dept.name,
                        'Limit': dept.budget_limit or 0,
                        'Zapotrzebowanie': dept_total,
                        'Różnica': dept_total - (dept.budget_limit or 0),
                        'Liczba pozycji': len(dept_entries)
                    })
                pd.DataFrame(dept_data).to_excel(writer, sheet_name='Departamenty', index=False)
            
            for sheet_name in writer.sheets:
                worksheet = writer.sheets[sheet_name]
                for column_cells in worksheet.columns:
                    max_length = 0
                    column_letter = column_cells[0].column_letter
                    for cell in column_cells:
                        try:
                            if cell.value:
                                max_length = max(max_length, len(str(cell.value)))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
        
        output.seek(0)
        return output
    
    def _get_global_limit(self, year: int) -> float:
        limit = self.db.query(GlobalLimit).filter(GlobalLimit.year == year).first()
        return limit.total_limit if limit else 0
    
    def export_limit_letter_to_docx(self, dept_code: str, year: int = 2025,
                                    new_limit: float = None) -> BytesIO:
        dept = self.db.query(Department).filter(Department.code == dept_code).first()
        if not dept:
            raise ValueError(f"Department {dept_code} not found")
        
        amount_field = getattr(BudgetEntry, f"kwota_{year}")
        entries = self.db.query(BudgetEntry).filter(
            BudgetEntry.department_id == dept.id,
            amount_field > 0
        ).all()
        
        dept_total = sum(getattr(e, f"kwota_{year}") or 0 for e in entries)
        assigned_limit = new_limit or dept.budget_limit or 0
        variance = dept_total - assigned_limit
        
        doc = Document()
        
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        header = doc.add_paragraph()
        header.add_run('Biuro Budżetowo-Finansowe\n').bold = True
        header.add_run('Ministerstwo Cyfryzacji\n')
        header.add_run(f'Warszawa, {date.today().strftime("%d.%m.%Y")}')
        
        ref = doc.add_paragraph()
        ref.add_run(f'Nr ref.: BBF-{date.today().year}/{dept_code}/{date.today().month:02d}')
        ref.paragraph_format.space_after = Pt(24)
        
        recipient = doc.add_paragraph()
        recipient.add_run(f'Pan/Pani Dyrektor\n').bold = True
        recipient.add_run(dept.name)
        recipient.paragraph_format.space_after = Pt(24)
        
        title = doc.add_paragraph()
        title_run = title.add_run(f'ZAWIADOMIENIE O LIMICIE WYDATKÓW NA ROK {year}')
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(18)
        
        doc.add_paragraph('Szanowny Panie Dyrektorze,')
        
        body1 = doc.add_paragraph()
        body1.add_run(
            f'Uprzejmie informuję, że zgodnie z pismem Ministerstwa Finansów, '
            f'limit wydatków dla {dept.name} na rok {year} został ustalony '
            f'na poziomie '
        )
        body1.add_run(f'{assigned_limit:,.0f} tys. PLN').bold = True
        body1.add_run('.')
        
        if variance > 0:
            body2 = doc.add_paragraph()
            body2.add_run(
                f'Zgłoszone przez Państwa zapotrzebowanie w wysokości '
            )
            body2.add_run(f'{dept_total:,.0f} tys. PLN').bold = True
            body2.add_run(' przekracza przyznany limit o ')
            body2.add_run(f'{variance:,.0f} tys. PLN').bold = True
            if assigned_limit > 0:
                body2.add_run(f' ({variance/assigned_limit*100:.1f}%).')
            else:
                body2.add_run('.')
            
            doc.add_paragraph(
                'Proszę o dokonanie przeglądu zgłoszonych potrzeb i wskazanie '
                'zadań do odroczenia lub redukcji w terminie do 7 dni roboczych.'
            )
        else:
            body2 = doc.add_paragraph()
            body2.add_run(
                f'Zgłoszone przez Państwa zapotrzebowanie w wysokości '
            )
            body2.add_run(f'{dept_total:,.0f} tys. PLN').bold = True
            body2.add_run(' mieści się w przyznanym limicie. Pozostała rezerwa wynosi ')
            body2.add_run(f'{abs(variance):,.0f} tys. PLN').bold = True
            body2.add_run('.')
        
        obligatory = [e for e in entries if e.is_obligatory]
        if obligatory:
            obligatory_sum = sum(getattr(e, f"kwota_{year}") or 0 for e in obligatory)
            oblig_para = doc.add_paragraph()
            oblig_para.add_run(
                'Przypominam, że zadania obligatoryjne (wynikające z przepisów prawa) '
                'w wysokości '
            )
            oblig_para.add_run(f'{obligatory_sum:,.0f} tys. PLN').bold = True
            oblig_para.add_run(f' ({len(obligatory)} pozycji) muszą zostać zabezpieczone w pierwszej kolejności.')
        
        doc.add_paragraph(
            f'Zgodnie z harmonogramem prac nad budżetem, ostateczne uzgodnienia '
            f'powinny zostać dokonane do końca sierpnia {year-1} r.'
        )
        
        doc.add_paragraph()
        table_title = doc.add_paragraph()
        table_title.add_run('Załącznik: Zestawienie pozycji budżetowych').bold = True
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Lp.'
        hdr_cells[1].text = 'Nazwa zadania'
        hdr_cells[2].text = 'Paragraf'
        hdr_cells[3].text = 'Kwota (tys. PLN)'
        hdr_cells[4].text = 'Priorytet'
        
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
        
        for i, entry in enumerate(sorted(entries, key=lambda e: getattr(e, f"kwota_{year}") or 0, reverse=True)[:15], 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = (entry.nazwa_zadania or entry.opis_projektu or '')[:50]
            row_cells[2].text = str(entry.paragraf or '')
            row_cells[3].text = f'{getattr(entry, f"kwota_{year}") or 0:,.0f}'
            row_cells[4].text = entry.priority if entry.priority else 'medium'
        
        doc.add_paragraph()
        closing = doc.add_paragraph('Z poważaniem,')
        closing.paragraph_format.space_after = Pt(36)
        
        signature = doc.add_paragraph()
        signature.add_run('Dyrektor\nBiura Budżetowo-Finansowego').bold = True
        
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    
    def export_summary_report_to_docx(self, year: int = 2025) -> BytesIO:
        global_limit = self.db.query(GlobalLimit).filter(GlobalLimit.year == year).first()
        
        amount_field = getattr(BudgetEntry, f"kwota_{year}")
        entries = self.db.query(BudgetEntry).filter(amount_field > 0).all()
        total = sum(getattr(e, f"kwota_{year}") or 0 for e in entries)
        
        doc = Document()
        
        title = doc.add_heading(f'INFORMACJA ZBIORCZA O BUDŻECIE NA ROK {year}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        meta = doc.add_paragraph()
        meta.add_run(f'Data wygenerowania: {datetime.now().strftime("%d.%m.%Y %H:%M")}\n')
        meta.add_run(f'Ministerstwo Cyfryzacji - Część budżetowa 27')
        meta.paragraph_format.space_after = Pt(24)
        
        doc.add_heading('1. Podsumowanie wykonawcze', level=1)
        
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_data = [
            ('Limit globalny MF', f'{global_limit.total_limit if global_limit else 0:,.0f} tys. PLN'),
            ('Łączne zapotrzebowanie', f'{total:,.0f} tys. PLN'),
            ('Różnica', f'{total - (global_limit.total_limit if global_limit else 0):,.0f} tys. PLN'),
            ('Status', 'PRZEKROCZENIE LIMITU' if total > (global_limit.total_limit if global_limit else 0) else 'W LIMICIE')
        ]
        
        for i, (label, value) in enumerate(summary_data):
            summary_table.rows[i].cells[0].text = label
            summary_table.rows[i].cells[1].text = value
        
        doc.add_heading('2. Podział na departamenty', level=1)
        
        dept_table = doc.add_table(rows=1, cols=5)
        dept_table.style = 'Table Grid'
        
        hdr = dept_table.rows[0].cells
        hdr[0].text = 'Departament'
        hdr[1].text = 'Limit'
        hdr[2].text = 'Zapotrzebowanie'
        hdr[3].text = 'Różnica'
        hdr[4].text = 'Pozycji'
        
        for dept in self.db.query(Department).all():
            dept_entries = [e for e in entries if e.department_id == dept.id]
            if not dept_entries:
                continue
            dept_total = sum(getattr(e, f"kwota_{year}") or 0 for e in dept_entries)
            
            row = dept_table.add_row().cells
            row[0].text = dept.code
            row[1].text = f'{dept.budget_limit or 0:,.0f}'
            row[2].text = f'{dept_total:,.0f}'
            row[3].text = f'{dept_total - (dept.budget_limit or 0):,.0f}'
            row[4].text = str(len(dept_entries))
        
        doc.add_heading('3. Podział wg priorytetów', level=1)
        
        prio_table = doc.add_table(rows=1, cols=3)
        prio_table.style = 'Table Grid'
        prio_table.rows[0].cells[0].text = 'Priorytet'
        prio_table.rows[0].cells[1].text = 'Liczba pozycji'
        prio_table.rows[0].cells[2].text = 'Kwota (tys. PLN)'
        
        for priority in PriorityLevel:
            prio_entries = [e for e in entries if e.priority == priority]
            if not prio_entries:
                continue
            prio_total = sum(getattr(e, f"kwota_{year}") or 0 for e in prio_entries)
            
            row = prio_table.add_row().cells
            row[0].text = priority.value.upper()
            row[1].text = str(len(prio_entries))
            row[2].text = f'{prio_total:,.0f}'
        
        doc.add_heading('4. Rekomendacje', level=1)
        
        variance = total - (global_limit.total_limit if global_limit else 0)
        if variance > 0:
            doc.add_paragraph(
                f'⚠️ Budżet przekracza limit o {variance:,.0f} tys. PLN. '
                f'Wymagana redukcja lub negocjacja zwiększenia limitu z Ministerstwem Finansów.',
                style='List Bullet'
            )
        else:
            doc.add_paragraph(
                f'✓ Budżet mieści się w limicie z rezerwą {abs(variance):,.0f} tys. PLN.',
                style='List Bullet'
            )
        
        obligatory_total = sum(getattr(e, f"kwota_{year}") or 0 for e in entries if e.is_obligatory)
        doc.add_paragraph(
            f'🔒 Zadania obligatoryjne wymagające zabezpieczenia: {obligatory_total:,.0f} tys. PLN',
            style='List Bullet'
        )
        
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output
